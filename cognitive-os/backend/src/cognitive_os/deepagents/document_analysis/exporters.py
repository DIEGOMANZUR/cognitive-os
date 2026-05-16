from __future__ import annotations

import csv
import io
import json
from pathlib import Path

from cognitive_os.deepagents.document_analysis.schemas import DocumentAnalysisResult


class DocumentAnalysisExporter:
    def export_json(self, result: DocumentAnalysisResult, output_dir: Path) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        path = output_dir / "result.json"
        path.write_text(
            json.dumps(result.model_dump(mode="json"), indent=2, ensure_ascii=False),
            encoding="utf-8",
        )
        return path

    def export_markdown(self, result: DocumentAnalysisResult, output_dir: Path) -> Path:
        output_dir.mkdir(parents=True, exist_ok=True)
        path = output_dir / "report.md"
        path.write_text(_markdown(result), encoding="utf-8")
        return path

    def export_csvs(self, result: DocumentAnalysisResult, output_dir: Path) -> list[Path]:
        """Emit evidence_matrix.csv, timeline.csv and contradictions.csv per spec."""
        output_dir.mkdir(parents=True, exist_ok=True)
        paths: list[Path] = []
        matrix_path = output_dir / "evidence_matrix.csv"
        matrix_path.write_text(_evidence_matrix_csv(result), encoding="utf-8")
        paths.append(matrix_path)
        timeline_path = output_dir / "timeline.csv"
        timeline_path.write_text(_timeline_csv(result), encoding="utf-8")
        paths.append(timeline_path)
        contradictions_path = output_dir / "contradictions.csv"
        contradictions_path.write_text(_contradictions_csv(result), encoding="utf-8")
        paths.append(contradictions_path)
        return paths

    def export_docx(self, result: DocumentAnalysisResult, output_dir: Path) -> Path | None:
        try:
            from docx import Document as DocxDocument
        except Exception:
            return None
        output_dir.mkdir(parents=True, exist_ok=True)
        path = output_dir / "report.docx"
        document = DocxDocument()
        document.add_heading("Document Analysis Report", 0)
        document.add_heading("Resumen ejecutivo", level=1)
        document.add_paragraph(result.executive_summary)
        document.add_heading("Matriz hecho/evidencia/cita", level=1)
        for row in result.evidence_matrix:
            document.add_paragraph(f"{row.claim_id}: {row.claim} [{row.strength}]")
        document.add_heading("Linea de tiempo", level=1)
        for event in result.timeline:
            document.add_paragraph(f"{event.date_text}: {event.event_summary}")
        document.add_heading("Contradicciones", level=1)
        for contradiction in result.contradictions:
            document.add_paragraph(f"{contradiction.topic}: {contradiction.explanation}")
        document.add_heading("Incertidumbres", level=1)
        for note in result.uncertainty_notes:
            document.add_paragraph(note)
        document.save(str(path))
        return path


def _markdown(result: DocumentAnalysisResult) -> str:
    sections = [
        "# Document Analysis Report",
        "## Resumen ejecutivo",
        result.executive_summary,
        "## Matriz hecho/evidencia/cita",
        _matrix_markdown(result),
        "## Línea de tiempo",
        _timeline_markdown(result),
        "## Contradicciones",
        _contradictions_markdown(result),
        "## Vacíos probatorios",
        _missing_markdown(result),
        "## Incertidumbres",
        "\n".join(f"- {note}" for note in result.uncertainty_notes) or "- Sin notas.",
        "## Anexos/citas",
        _citations_markdown(result),
    ]
    return "\n\n".join(sections) + "\n"


def _matrix_markdown(result: DocumentAnalysisResult) -> str:
    header = (
        "| claim_id | tipo | afirmación | fuerza | citas | notas |\n"
        "| --- | --- | --- | --- | --- | --- |"
    )
    rows = []
    for row in result.evidence_matrix:
        citations = ", ".join(_citation_text(citation) for citation in row.supporting_evidence)
        rows.append(
            f"| {row.claim_id} | {row.claim_type} | {row.claim} | {row.strength} | "
            f"{citations or 'sin cita'} | {row.notes} |"
        )
    return "\n".join([header, *rows]) if rows else "Sin matriz."


def _timeline_markdown(result: DocumentAnalysisResult) -> str:
    if not result.timeline:
        return "Sin timeline."
    return "\n".join(
        f"- {event.date_text} ({event.date_certainty}): {event.event_summary}"
        for event in result.timeline
    )


def _contradictions_markdown(result: DocumentAnalysisResult) -> str:
    if not result.contradictions:
        return "Sin contradicciones detectadas."
    return "\n".join(
        f"- {item.topic}: {item.statement_a} / {item.statement_b}. {item.explanation}"
        for item in result.contradictions
    )


def _missing_markdown(result: DocumentAnalysisResult) -> str:
    if not result.missing_evidence:
        return "Sin vacíos probatorios registrados."
    return "\n".join(
        f"- {item.expected_evidence}: {item.why_it_matters} [{item.status}]"
        for item in result.missing_evidence
    )


def _citations_markdown(result: DocumentAnalysisResult) -> str:
    if not result.citations:
        return "Sin citas."
    return "\n".join(f"- {_citation_text(citation)}" for citation in result.citations)


def _citation_text(citation: object) -> str:
    doc_id = getattr(citation, "doc_id", "")
    page_start = getattr(citation, "page_start", "")
    page_end = getattr(citation, "page_end", "")
    chunk_id = getattr(citation, "chunk_id", None)
    suffix = f", chunk_id={chunk_id}" if chunk_id else ""
    return f"doc_id={doc_id}, pages={page_start}-{page_end}{suffix}"


def _csv_dump(header: list[str], rows: list[list[str]]) -> str:
    buffer = io.StringIO()
    writer = csv.writer(buffer, lineterminator="\n")
    writer.writerow(header)
    writer.writerows(rows)
    return buffer.getvalue()


def _evidence_matrix_csv(result: DocumentAnalysisResult) -> str:
    header = [
        "claim_id",
        "claim_type",
        "claim",
        "strength",
        "supporting_evidence",
        "opposing_evidence",
        "needs_human_review",
        "notes",
    ]
    rows: list[list[str]] = []
    for row in result.evidence_matrix:
        rows.append(
            [
                row.claim_id,
                row.claim_type,
                row.claim,
                row.strength,
                "; ".join(_citation_text(citation) for citation in row.supporting_evidence),
                "; ".join(_citation_text(citation) for citation in row.opposing_evidence),
                "true" if row.needs_human_review else "false",
                row.notes,
            ]
        )
    return _csv_dump(header, rows)


def _timeline_csv(result: DocumentAnalysisResult) -> str:
    header = [
        "event_id",
        "date_text",
        "normalized_date",
        "date_certainty",
        "event_summary",
        "involved_entities",
        "citations",
        "contradictions",
        "notes",
    ]
    rows: list[list[str]] = []
    for event in result.timeline:
        rows.append(
            [
                event.event_id,
                event.date_text,
                event.normalized_date.isoformat() if event.normalized_date else "",
                event.date_certainty,
                event.event_summary,
                "; ".join(event.involved_entities),
                "; ".join(_citation_text(citation) for citation in event.citations),
                "; ".join(event.contradictions),
                event.notes,
            ]
        )
    return _csv_dump(header, rows)


def _contradictions_csv(result: DocumentAnalysisResult) -> str:
    header = [
        "contradiction_id",
        "topic",
        "contradiction_type",
        "severity",
        "statement_a",
        "citation_a",
        "statement_b",
        "citation_b",
        "needs_human_review",
        "explanation",
    ]
    rows: list[list[str]] = []
    for contradiction in result.contradictions:
        rows.append(
            [
                contradiction.contradiction_id,
                contradiction.topic,
                contradiction.contradiction_type,
                contradiction.severity,
                contradiction.statement_a,
                _citation_text(contradiction.citation_a),
                contradiction.statement_b,
                _citation_text(contradiction.citation_b),
                "true" if contradiction.needs_human_review else "false",
                contradiction.explanation,
            ]
        )
    return _csv_dump(header, rows)
